"""Unified content fetcher — routes by content type, caches in MongoDB.

Single entry point for fetching content from any source registered in the
Table of Content database: Notion Pages, Notion Sites, Google Docs, Sheets.

Architecture:
  1. Query ToC database → get all rows
  2. Route each row to the right fetcher based on Content type
  3. Cache fetched content in MongoDB (notion_page_cache) with TTL
  4. Skip unchanged content via content hash comparison

Fetch chains per content type:
  Notion Page  → extract page_id → MCP fetch_page() → REST API fallback
  Notion Site  → extract page_id from URL slug → MCP fetch_page() → HTTP fetch public URL
  Google Docx  → extract doc_id + tab_id → Docs API per-tab → Drive API export fallback
  Sheets       → Drive API export as CSV → skip with warning
"""
from __future__ import annotations

import hashlib
import logging
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Set

import httpx

logger = logging.getLogger(__name__)

# Cache TTL: 24 hours (pages older than this are re-fetched)
CACHE_TTL_SECONDS = 24 * 60 * 60
MIN_WORDS = 40


# ── Page ID extraction ───────────────────────────────────────────────────────

def extract_notion_page_id(url: str, page_id_column: str = "") -> Optional[str]:
    """Extract a 32-char hex page ID from any Notion URL or column value.

    Handles:
      - notion.so/workspace/Title-{32hex}
      - altcarbon.notion.site/slug-{32hex}
      - notion.site/slug-{32hex}?query
      - Raw 32-hex ID in the Page ID column
      - UUID-formatted IDs (with dashes)
    """
    # Try the explicit Page ID column first
    if page_id_column:
        clean = page_id_column.strip().replace("-", "")
        match = re.search(r'([a-f0-9]{32})', clean)
        if match:
            return match.group(1)

    if not url:
        return None

    # Check if URL is a Notion URL (workspace or published site)
    if "notion.so" not in url and "notion.site" not in url:
        return None

    # Strip query params and trailing slash
    clean_path = url.split("?")[0].split("#")[0].rstrip("/")
    last_segment = clean_path.split("/")[-1]

    # The page ID is always the last 32 hex chars in the slug
    # e.g., "about-us-24e50d0ec20e80d2803fc81..." or "Title-1f350d0ec20e803988dbde4..."
    segment_clean = last_segment.replace("-", "")
    match = re.search(r'([a-f0-9]{32})$', segment_clean)
    if match:
        return match.group(1)

    # Also try with dashes (UUID format embedded in URL)
    uuid_match = re.search(
        r'([a-f0-9]{8}-?[a-f0-9]{4}-?[a-f0-9]{4}-?[a-f0-9]{4}-?[a-f0-9]{12})',
        url,
    )
    if uuid_match:
        return uuid_match.group(1).replace("-", "")

    return None


def extract_drive_doc_id(url: str) -> Optional[str]:
    """Extract Google Drive document ID from a URL."""
    if not url:
        return None
    match = re.search(r'docs\.google\.com/document/d/([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)
    match = re.search(r'docs\.google\.com/spreadsheets/d/([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)
    return None


def extract_tab_id(url: str) -> Optional[str]:
    """Extract Google Docs tab ID from a URL query param like ?tab=t.{tabId}.

    Returns the full tab ID including 't.' prefix (e.g., 't.0', 't.62emwh25ib2z')
    to match the Google Docs API tabProperties.tabId format.
    """
    if not url:
        return None
    match = re.search(r'[?&]tab=(t\.[a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)
    return None


# ── MongoDB page cache ───────────────────────────────────────────────────────

def _page_cache():
    from backend.db.mongo import get_db
    return get_db()["notion_page_cache"]


async def _get_cached(source_id: str) -> Optional[Dict]:
    """Get cached page content if fresh (within TTL)."""
    doc = await _page_cache().find_one({"source_id": source_id})
    if not doc:
        return None

    cached_at = doc.get("cached_at")
    if not cached_at:
        return None

    # Check TTL
    if isinstance(cached_at, str):
        cached_at = datetime.fromisoformat(cached_at)
    if not cached_at.tzinfo:
        cached_at = cached_at.replace(tzinfo=timezone.utc)

    age = (datetime.now(timezone.utc) - cached_at).total_seconds()
    if age > CACHE_TTL_SECONDS:
        return None

    return doc


async def _set_cached(source_id: str, title: str, text: str, source: str, content_type: str):
    """Cache fetched page content in MongoDB."""
    await _page_cache().update_one(
        {"source_id": source_id},
        {"$set": {
            "source_id": source_id,
            "title": title,
            "text": text,
            "source": source,
            "content_type": content_type,
            "content_hash": hashlib.sha256(text.encode()).hexdigest(),
            "chars": len(text),
            "cached_at": datetime.now(timezone.utc),
        }},
        upsert=True,
    )


# ── Content fetchers per type ────────────────────────────────────────────────

# Max recursion depth for nested blocks (toggles, callouts, child pages)
_MAX_BLOCK_DEPTH = 3


async def _extract_blocks_recursive(
    client: httpx.AsyncClient,
    block_id: str,
    headers: Dict,
    depth: int = 0,
) -> List[str]:
    """Recursively extract text from Notion blocks, including children.

    Handles toggle blocks, callouts, child pages, columns, synced blocks, etc.
    """
    if depth > _MAX_BLOCK_DEPTH:
        return []

    lines: List[str] = []
    cursor: Optional[str] = None

    while True:
        params: Dict[str, Any] = {"page_size": 100}
        if cursor:
            params["start_cursor"] = cursor

        r = await client.get(
            f"https://api.notion.com/v1/blocks/{block_id}/children",
            headers=headers,
            params=params,
        )
        if r.status_code != 200:
            break

        data = r.json()
        blocks = data.get("results", [])

        for block in blocks:
            btype = block.get("type", "")
            bid = block.get("id", "")
            has_children = block.get("has_children", False)
            content = block.get(btype, {})

            # Extract text from rich_text (paragraphs, headings, lists, toggles, callouts, quotes)
            rich = content.get("rich_text", [])
            text = "".join(rt.get("plain_text", "") for rt in rich)

            # Special block types
            if btype == "child_page":
                page_title = content.get("title", "")
                if page_title:
                    lines.append(f"\n## {page_title}\n")
            elif btype == "child_database":
                db_title = content.get("title", "")
                if db_title:
                    lines.append(f"\n[Database: {db_title}]\n")
            elif btype == "image":
                caption = "".join(rt.get("plain_text", "") for rt in content.get("caption", []))
                img_url = ""
                if content.get("type") == "file":
                    img_url = content.get("file", {}).get("url", "")
                elif content.get("type") == "external":
                    img_url = content.get("external", {}).get("url", "")
                if caption:
                    lines.append(f"[Image: {caption}]")
                elif img_url:
                    lines.append(f"[Image: {img_url[:100]}]")
            elif btype == "bookmark":
                bm_url = content.get("url", "")
                caption = "".join(rt.get("plain_text", "") for rt in content.get("caption", []))
                if caption:
                    lines.append(f"[Bookmark: {caption} — {bm_url}]")
                elif bm_url:
                    lines.append(f"[Bookmark: {bm_url}]")
            elif btype == "table":
                # Tables: fetch children rows
                pass  # handled by recursion below
            elif btype in ("heading_1", "heading_2", "heading_3"):
                level = {"heading_1": "#", "heading_2": "##", "heading_3": "###"}[btype]
                if text.strip():
                    lines.append(f"\n{level} {text.strip()}\n")
                    text = ""  # already added
            elif btype == "table_row":
                cells = content.get("cells", [])
                row_text = " | ".join(
                    "".join(rt.get("plain_text", "") for rt in cell)
                    for cell in cells
                )
                if row_text.strip():
                    lines.append(row_text)
                    text = ""
            elif btype in ("bulleted_list_item", "numbered_list_item"):
                if text.strip():
                    lines.append(f"• {text.strip()}")
                    text = ""
            elif btype == "to_do":
                checked = content.get("checked", False)
                mark = "x" if checked else " "
                if text.strip():
                    lines.append(f"[{mark}] {text.strip()}")
                    text = ""
            elif btype == "code":
                lang = content.get("language", "")
                if text.strip():
                    lines.append(f"```{lang}\n{text.strip()}\n```")
                    text = ""
            elif btype == "equation":
                expr = content.get("expression", "")
                if expr:
                    lines.append(f"$$ {expr} $$")
            elif btype == "divider":
                lines.append("---")

            # Add remaining text (paragraphs, callouts, toggles, quotes)
            if text.strip():
                lines.append(text.strip())

            # Recurse into children (toggles, callouts, columns, synced blocks, child pages)
            if has_children and btype != "child_database":
                child_lines = await _extract_blocks_recursive(
                    client, bid, headers, depth + 1
                )
                lines.extend(child_lines)

        cursor = data.get("next_cursor")
        if not cursor:
            break

    return lines


async def fetch_notion_page(page_id: str, title: str = "") -> Optional[str]:
    """Fetch a Notion page by ID. MCP first, REST API with recursive block fetch fallback."""
    # Try MCP (already handles recursive fetch internally)
    try:
        from backend.integrations.notion_mcp import notion_mcp
        if notion_mcp.connected:
            text = await notion_mcp.fetch_page(page_id)
            if text and len(text.split()) >= MIN_WORDS:
                logger.debug("Fetched '%s' via MCP (%d chars)", title[:40], len(text))
                return text
    except Exception as e:
        logger.warning("Notion MCP fetch failed for %s: %s", page_id, e)

    # Fallback: REST API with recursive block extraction
    try:
        from backend.config.settings import get_settings
        token = get_settings().notion_token
        if not token:
            return None

        headers = {
            "Authorization": f"Bearer {token}",
            "Notion-Version": "2022-06-28",
        }
        async with httpx.AsyncClient(timeout=60.0) as client:
            lines = await _extract_blocks_recursive(client, page_id, headers, depth=0)

        result = "\n".join(lines)
        if len(result.split()) >= MIN_WORDS:
            logger.debug("Fetched '%s' via REST API recursive (%d chars)", title[:40], len(result))
            return result
    except Exception as e:
        logger.warning("REST API fetch failed for %s: %s", page_id, e)

    return None


async def fetch_with_cloudflare(url: str) -> Optional[str]:
    """Fetch a page via Cloudflare Browser Rendering (renders JS, returns markdown)."""
    from backend.config.settings import get_settings
    s = get_settings()
    account_id = getattr(s, "cloudflare_account_id", "")
    api_token = getattr(s, "cloudflare_browser_token", "")

    if not account_id or not api_token:
        return None

    cf_url = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/browser-rendering/markdown"
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            r = await client.post(
                cf_url,
                headers={
                    "Authorization": f"Bearer {api_token}",
                    "Content-Type": "application/json",
                },
                json={"url": url.strip()},
            )
            r.raise_for_status()
            data = r.json()
            text = (data.get("result") or "").strip()[:80_000]
            if text and len(text.split()) >= MIN_WORDS:
                logger.debug("Fetched '%s' via Cloudflare BR (%d chars)", url[:60], len(text))
                return text
    except Exception as e:
        logger.debug("Cloudflare BR fetch failed for %s: %s", url[:60], e)

    return None


async def fetch_with_browser(url: str) -> Optional[str]:
    """Fetch a page via agent-browser headless Chromium (renders JS fully)."""
    try:
        from backend.utils.browser import browser_fetch, is_available
        if not is_available():
            return None
        text = await browser_fetch(url, timeout=45.0)
        if text and len(text.split()) >= MIN_WORDS:
            logger.info("Fetched '%s' via agent-browser (%d chars)", url[:60], len(text))
            return text
    except Exception as e:
        logger.debug("agent-browser fetch failed for %s: %s", url[:60], e)
    return None


async def fetch_notion_internal_api(page_id: str, site_domain: str = "altcarbon.notion.site", title: str = "") -> Optional[str]:
    """Fetch a public Notion page via the internal API (no auth needed).

    Works for published notion.site pages that aren't shared with the API integration.
    Uses loadPageChunk to get block content.
    """
    pid = f"{page_id[:8]}-{page_id[8:12]}-{page_id[12:16]}-{page_id[16:20]}-{page_id[20:]}"

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            all_blocks: Dict[str, Any] = {}
            cursor: Dict = {"stack": []}
            chunk = 0

            # Fetch up to 3 chunks (300 blocks)
            while chunk < 3:
                r = await client.post(
                    f"https://{site_domain}/api/v3/loadPageChunk",
                    json={
                        "page": {"id": pid},
                        "limit": 100,
                        "cursor": cursor,
                        "chunkNumber": chunk,
                        "verticalColumns": False,
                    },
                    headers={"Content-Type": "application/json"},
                )
                if r.status_code != 200:
                    break

                data = r.json()
                blocks = data.get("recordMap", {}).get("block", {})
                all_blocks.update(blocks)

                # Check if there are more chunks
                new_cursor = data.get("cursor", {})
                if not new_cursor.get("stack"):
                    break
                cursor = new_cursor
                chunk += 1

            if not all_blocks:
                return None

            # Extract text from blocks
            lines: List[str] = []
            for bid, bdata in all_blocks.items():
                value = bdata.get("value", {})
                btype = value.get("type", "")
                props = value.get("properties", {})

                title_arr = props.get("title", [])
                text = "".join(
                    t[0] if isinstance(t, list) and t else "" for t in title_arr
                )

                if text.strip():
                    if btype in ("header", "sub_header", "sub_sub_header"):
                        level = 1 + ["header", "sub_header", "sub_sub_header"].index(btype)
                        lines.append(f"\n{'#' * level} {text.strip()}\n")
                    elif btype == "bulleted_list":
                        lines.append(f"• {text.strip()}")
                    elif btype == "numbered_list":
                        lines.append(f"- {text.strip()}")
                    elif btype == "callout":
                        lines.append(f"> {text.strip()}")
                    elif btype == "quote":
                        lines.append(f"> {text.strip()}")
                    elif btype == "code":
                        lines.append(f"```\n{text.strip()}\n```")
                    elif btype == "image":
                        caption = "".join(
                            t[0] if isinstance(t, list) and t else "" for t in props.get("caption", [])
                        )
                        if caption.strip():
                            lines.append(f"[Image: {caption.strip()}]")
                    else:
                        lines.append(text.strip())

            result = "\n".join(lines)
            if len(result.split()) >= MIN_WORDS:
                logger.info("Fetched '%s' via internal API (%d chars, %d blocks)", title[:40], len(result), len(all_blocks))
                return result

    except Exception as e:
        logger.debug("Internal API fetch failed for %s: %s", page_id, e)

    return None


async def _resolve_notion_site_page_id(url: str) -> Optional[str]:
    """Resolve a Notion page ID from a public notion.site URL.

    Fetches the HTML and extracts the page ID from embedded metadata
    (Notion injects it in the initial HTML even before JS renders).
    """
    try:
        async with httpx.AsyncClient(
            timeout=20.0, follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0"},
        ) as client:
            r = await client.get(url)
            if r.status_code != 200:
                return None
            html = r.text
            # Notion embeds pageId in the initial HTML response
            for pat in (
                r'"pageId"\s*:\s*"([a-f0-9-]{36})"',
                r'"block_id"\s*:\s*"([a-f0-9-]{36})"',
                r'data-page-id="([a-f0-9-]{36})"',
            ):
                m = re.search(pat, html)
                if m:
                    resolved = m.group(1).replace("-", "")
                    logger.debug("Resolved page ID for %s → %s", url[:60], resolved)
                    return resolved
    except Exception as e:
        logger.debug("Failed to resolve page ID from %s: %s", url[:60], e)
    return None


async def fetch_notion_site(url: str, page_id: Optional[str], title: str = "") -> Optional[str]:
    """Fetch a Notion Site page.

    Fallback chain:
      1. MCP/REST API with page_id (if available)
      2. Resolve page_id from HTML metadata → REST API recursive fetch
      3. agent-browser headless Chromium (full JS rendering)
      4. Cloudflare Browser Rendering
      5. HTTP fetch public URL (basic HTML extraction)
    """
    # If we have a page ID, use the same path as Notion Page
    if page_id:
        result = await fetch_notion_page(page_id, title=title)
        if result:
            return result

    # Resolve page ID from the public URL HTML
    resolved_id: Optional[str] = None
    if url and not page_id:
        resolved_id = await _resolve_notion_site_page_id(url)
        if resolved_id:
            # Try REST API first (works if page is shared with integration)
            result = await fetch_notion_page(resolved_id, title=title)
            if result:
                return result

    # Internal API (works for public pages without API integration sharing)
    effective_id = page_id or resolved_id
    if effective_id and url:
        # Extract domain from URL for the internal API
        import urllib.parse
        domain = urllib.parse.urlparse(url).netloc
        result = await fetch_notion_internal_api(effective_id, site_domain=domain, title=title)
        if result:
            return result

    # agent-browser: full headless Chromium rendering
    if url:
        result = await fetch_with_browser(url)
        if result:
            return result

    # Cloudflare Browser Rendering
    if url:
        result = await fetch_with_cloudflare(url)
        if result:
            return result

    # Last resort: HTTP fetch (usually gets JS shell for Notion Sites)
    if url:
        try:
            async with httpx.AsyncClient(
                timeout=30.0,
                follow_redirects=True,
                headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"},
            ) as client:
                r = await client.get(url)
                if r.status_code == 200:
                    html = r.text
                    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
                    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
                    text = re.sub(r'<[^>]+>', ' ', html)
                    text = re.sub(r'\s+', ' ', text).strip()
                    if len(text.split()) >= MIN_WORDS:
                        logger.debug("Fetched '%s' via HTTP (%d chars)", title[:40], len(text))
                        return text
        except Exception as e:
            logger.debug("HTTP fetch failed for %s: %s", url, e)

    return None


def _build_google_creds(creds: Dict) -> Optional[Any]:
    """Build Google OAuth credentials from kwargs or settings."""
    refresh_token = creds.get("google_refresh_token", "")
    client_id = creds.get("google_client_id", "")
    client_secret = creds.get("google_client_secret", "")

    if not refresh_token:
        from backend.config.settings import get_settings
        s = get_settings()
        refresh_token = s.google_refresh_token
        client_id = client_id or s.google_client_id
        client_secret = client_secret or s.google_client_secret

    if not refresh_token:
        return None

    from google.oauth2.credentials import Credentials
    return Credentials(
        token=None,
        refresh_token=refresh_token,
        client_id=client_id,
        client_secret=client_secret,
        token_uri="https://oauth2.googleapis.com/token",
        scopes=["https://www.googleapis.com/auth/drive.readonly"],
    )


def _extract_tab_text(tab: Dict) -> str:
    """Extract plain text from a Google Docs API tab object.

    Walks the tab's document body content (structural elements → paragraphs →
    text runs) and concatenates all text.
    """
    body = tab.get("documentTab", {}).get("body", {})
    parts: List[str] = []
    for element in body.get("content", []):
        paragraph = element.get("paragraph")
        if not paragraph:
            # Handle tables
            table = element.get("table")
            if table:
                for row in table.get("tableRows", []):
                    for cell in row.get("tableCells", []):
                        for cell_el in cell.get("content", []):
                            p = cell_el.get("paragraph")
                            if p:
                                for run in p.get("elements", []):
                                    tr = run.get("textRun", {}).get("content", "")
                                    if tr.strip():
                                        parts.append(tr)
            continue
        for run in paragraph.get("elements", []):
            tr = run.get("textRun", {}).get("content", "")
            if tr:
                parts.append(tr)
    return "".join(parts).strip()


async def fetch_google_doc_tab(
    doc_id: str,
    tab_id: str,
    title: str = "",
    **creds,
) -> Optional[str]:
    """Fetch a specific tab from a Google Doc using the Docs API.

    Uses documents().get(includeTabsContent=True) to retrieve all tabs,
    then extracts the content for the matching tab_id.
    """
    import asyncio

    credential = _build_google_creds(creds)
    if not credential:
        logger.debug("No Google credentials — skipping Docs tab %s/%s", doc_id, tab_id)
        return None

    try:
        from googleapiclient.discovery import build

        service = await asyncio.to_thread(
            build, "docs", "v1", credentials=credential, cache_discovery=False
        )
        doc = await asyncio.to_thread(
            lambda: service.documents().get(
                documentId=doc_id, includeTabsContent=True
            ).execute()
        )

        tabs = doc.get("tabs", [])

        # Flatten nested tabs (Google Docs supports child tabs)
        def _flatten_tabs(tab_list: List[Dict]) -> List[Dict]:
            flat: List[Dict] = []
            for t in tab_list:
                flat.append(t)
                children = t.get("childTabs", [])
                if children:
                    flat.extend(_flatten_tabs(children))
            return flat

        all_tabs = _flatten_tabs(tabs)

        # Find the matching tab
        for tab in all_tabs:
            props = tab.get("tabProperties", {})
            if props.get("tabId") == tab_id:
                text = _extract_tab_text(tab)
                if text and len(text.split()) >= MIN_WORDS:
                    logger.debug(
                        "Fetched tab '%s' (%s) from '%s' via Docs API (%d chars)",
                        props.get("title", ""), tab_id, title[:40], len(text),
                    )
                    return text
                else:
                    logger.debug(
                        "Tab '%s' (%s) in '%s' has insufficient text (%d words)",
                        props.get("title", ""), tab_id, title[:40],
                        len(text.split()) if text else 0,
                    )
                    return None

        # Tab not found — log available tabs for debugging
        available = [
            f"{t.get('tabProperties', {}).get('title', '?')} ({t.get('tabProperties', {}).get('tabId', '?')})"
            for t in all_tabs
        ]
        logger.warning(
            "Tab ID '%s' not found in doc '%s'. Available tabs: %s",
            tab_id, title, ", ".join(available),
        )
    except Exception as e:
        logger.warning("Docs API tab fetch failed for %s tab %s ('%s'): %s", doc_id, tab_id, title, e)

    return None


def _parse_docx_bytes(data: bytes) -> str:
    """Parse a .docx file from bytes, extracting paragraphs and table content."""
    import io
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    parts: List[str] = []

    # Paragraphs (section headings, body text)
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tables (where most structured content lives in Puro-style PDDs)
    for table in doc.tables:
        for row in table.rows:
            cells: List[str] = []
            seen: Set[str] = set()
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in seen:
                    cells.append(t)
                    seen.add(t)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _download_drive_file(service: Any, file_id: str) -> bytes:
    """Download a file from Google Drive via get_media(). Synchronous — call in to_thread."""
    import io as _io
    from googleapiclient.http import MediaIoBaseDownload

    request = service.files().get_media(fileId=file_id)
    fh = _io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return fh.getvalue()


async def fetch_google_doc(doc_id: str, title: str = "", tab_id: str = "", **creds) -> Optional[str]:
    """Fetch a Google Doc by ID.

    Handles three cases:
      1. Tab-based: uses Docs API per-tab fetch
      2. Native Google Doc: uses Drive API export as text/plain
      3. Uploaded .docx: downloads raw file and parses with python-docx
    """
    import asyncio

    # Tab-aware fetch: use Docs API
    if tab_id:
        result = await fetch_google_doc_tab(doc_id, tab_id, title=title, **creds)
        if result:
            return result
        logger.debug("Tab fetch failed for %s/%s, falling back to Drive export", doc_id, tab_id)

    credential = _build_google_creds(creds)
    if not credential:
        logger.debug("No Google credentials — skipping Drive doc %s", doc_id)
        return None

    try:
        from googleapiclient.discovery import build

        service = await asyncio.to_thread(
            build, "drive", "v3", credentials=credential, cache_discovery=False
        )

        # Check mimeType to decide export vs download
        meta = await asyncio.to_thread(
            lambda: service.files().get(fileId=doc_id, fields="mimeType,name").execute()
        )
        mime = meta.get("mimeType", "")

        if mime == "application/vnd.google-apps.document":
            # Native Google Doc → export as plain text
            export = await asyncio.to_thread(
                lambda: service.files().export(
                    fileId=doc_id, mimeType="text/plain"
                ).execute()
            )
            text = export.decode("utf-8") if isinstance(export, bytes) else str(export)
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            # Uploaded .docx/.doc → download and parse
            raw = await asyncio.to_thread(
                lambda: _download_drive_file(service, doc_id)
            )
            text = await asyncio.to_thread(_parse_docx_bytes, raw)
        else:
            # Other type — try export as plain text
            export = await asyncio.to_thread(
                lambda: service.files().export(
                    fileId=doc_id, mimeType="text/plain"
                ).execute()
            )
            text = export.decode("utf-8") if isinstance(export, bytes) else str(export)

        if text and len(text.split()) >= MIN_WORDS:
            logger.debug("Fetched '%s' via Drive API (%d chars, mime=%s)", title[:40], len(text), mime)
            return text
    except Exception as e:
        logger.warning("Drive doc fetch failed for %s ('%s'): %s", doc_id, title, e)

    return None


async def fetch_google_sheet(doc_id: str, title: str = "", **creds) -> Optional[str]:
    """Fetch a Google Sheet by ID, export as CSV via Drive API."""
    import asyncio

    credential = _build_google_creds(creds)
    if not credential:
        return None

    try:
        from googleapiclient.discovery import build

        service = await asyncio.to_thread(
            build, "drive", "v3", credentials=credential, cache_discovery=False
        )
        export = await asyncio.to_thread(
            lambda: service.files().export(
                fileId=doc_id, mimeType="text/csv"
            ).execute()
        )
        text = export.decode("utf-8") if isinstance(export, bytes) else str(export)
        if len(text.split()) >= MIN_WORDS:
            logger.debug("Fetched sheet '%s' via Drive API (%d chars)", title[:40], len(text))
            return text
    except Exception as e:
        logger.debug("Sheet fetch failed for %s ('%s'): %s", doc_id, title, e)

    return None


# ── Main router: fetch from Table of Content ─────────────────────────────────

async def fetch_toc_row(
    row: Dict[str, Any],
    google_creds: Optional[Dict] = None,
    use_cache: bool = True,
) -> Optional[Dict]:
    """Fetch content for a single Table of Content row.

    Routes to the correct fetcher based on Content type.
    Returns a doc dict ready for chunking, or None if fetch failed.
    """
    from backend.integrations.notion_config import TOC_THEME_MAP

    props = row.get("properties", {})

    # ── Extract metadata ─────────────────────────────────────────────────
    doc_name_prop = props.get("Document Name", {})
    doc_name = "".join(
        t.get("plain_text", "") for t in doc_name_prop.get("title", [])
    ).strip()
    display_name = re.sub(r'\*\*|<[^>]+>', '', doc_name).strip() or "Untitled"

    content_type = (props.get("Content type", {}).get("select") or {}).get("name", "")
    url_field = props.get("URL", {}).get("url") or ""
    extra_url = props.get("Extra page url", {}).get("url") or ""

    # Content info: themes + main-source flag
    content_info = [
        opt.get("name", "") for opt in props.get("Content info", {}).get("multi_select", [])
    ]
    is_main_source = "Main-source" in content_info
    themes = [TOC_THEME_MAP[t] for t in content_info if t in TOC_THEME_MAP]

    # Notion Page ID column
    notion_page_id_raw = "".join(
        t.get("plain_text", "")
        for t in props.get("Notion Page ID", {}).get("rich_text", [])
    ).strip()

    # ── Resolve source ID ────────────────────────────────────────────────
    page_id = extract_notion_page_id(url_field, notion_page_id_raw)
    drive_id = extract_drive_doc_id(url_field)
    tab_id = extract_tab_id(url_field) or ""
    extra_drive_id = extract_drive_doc_id(extra_url)
    # For tab-based docs, use doc_id:tab_id as source_id so each tab caches separately
    source_id = page_id or (f"{drive_id}:{tab_id}" if drive_id and tab_id else drive_id) or ""

    if not source_id and not url_field:
        logger.debug("ToC: no fetchable source for '%s' (%s)", display_name, content_type)
        return None

    # ── Check cache ──────────────────────────────────────────────────────
    if use_cache and source_id:
        cached = await _get_cached(source_id)
        if cached:
            logger.debug("ToC: cache hit for '%s' (%d chars)", display_name, cached.get("chars", 0))
            return {
                "id": source_id,
                "title": display_name,
                "text": cached["text"],
                "source": cached.get("source", "notion"),
                "priority": "high" if is_main_source else "medium",
                "source_registry": "table_of_content",
                "focus_areas": themes,
                "content_type": content_type,
            }

    # ── Fetch by content type ────────────────────────────────────────────
    text = None
    source = "notion"
    creds = google_creds or {}

    if content_type == "Notion Page" and page_id:
        text = await fetch_notion_page(page_id, title=display_name)

    elif content_type == "Notion Site":
        # Notion Site: try page_id first, then public URL
        text = await fetch_notion_site(url_field, page_id, title=display_name)
        if not text and not page_id:
            # Last resort: search by title in MCP
            try:
                from backend.integrations.notion_mcp import notion_mcp
                if notion_mcp.connected:
                    results = await notion_mcp.search(display_name)
                    for sr in results:
                        sr_id = sr.get("id", "").replace("-", "")
                        if sr_id and len(sr_id) == 32:
                            text = await fetch_notion_page(sr_id, title=display_name)
                            if text:
                                source_id = sr_id
                                break
            except Exception as e:
                logger.debug("ToC: MCP search for '%s' failed: %s", display_name, e)

    elif content_type == "Google Docx" and drive_id:
        text = await fetch_google_doc(drive_id, title=display_name, tab_id=tab_id, **creds)
        source = "drive"

    elif content_type == "Sheets" and drive_id:
        text = await fetch_google_sheet(drive_id, title=display_name, **creds)
        source = "drive"

    else:
        logger.debug("ToC: unsupported content type '%s' for '%s'", content_type, display_name)

    # ── Fetch extra URL (secondary source) ───────────────────────────────
    extra_text = None
    if extra_drive_id and extra_drive_id != drive_id:
        extra_text = await fetch_google_doc(extra_drive_id, title=f"{display_name} (extra)", **creds)

    # ── Combine results ──────────────────────────────────────────────────
    if not text and not extra_text:
        logger.debug("ToC: no content for '%s' (%s)", display_name, content_type)
        return None

    combined = text or ""
    if extra_text:
        combined = f"{combined}\n\n---\n\n{extra_text}" if combined else extra_text

    # ── Cache the result ─────────────────────────────────────────────────
    if source_id:
        try:
            await _set_cached(source_id, display_name, combined, source, content_type)
        except Exception as e:
            logger.debug("ToC: cache write failed for '%s': %s", display_name, e)

    return {
        "id": source_id or url_field.rstrip("/").split("/")[-1],
        "title": display_name,
        "text": combined,
        "source": source,
        "priority": "high" if is_main_source else "medium",
        "source_registry": "table_of_content",
        "focus_areas": themes,
        "content_type": content_type,
    }


async def fetch_all_from_toc(
    use_cache: bool = True,
    google_creds: Optional[Dict] = None,
) -> List[Dict]:
    """Query the Table of Content database and fetch all content.

    This is the primary entry point for Company Brain knowledge sync.
    Returns a list of doc dicts ready for chunking.
    """
    from backend.integrations.notion_config import TABLE_OF_CONTENT_DS

    # ── Query ToC database ───────────────────────────────────────────────
    rows: List[Dict] = []

    # Try MCP first
    try:
        from backend.integrations.notion_mcp import notion_mcp
        if notion_mcp.connected:
            rows = await notion_mcp.query_data_source(TABLE_OF_CONTENT_DS, limit=100)
            if rows:
                logger.info("ToC: queried %d rows via MCP", len(rows))
    except Exception as e:
        logger.warning("ToC: MCP query failed, falling back to REST: %s", e)

    # Fallback: REST API
    if not rows:
        try:
            from backend.config.settings import get_settings
            token = get_settings().notion_token
            if not token:
                logger.warning("ToC: no NOTION_TOKEN — cannot query Table of Content")
                return []

            headers = {
                "Authorization": f"Bearer {token}",
                "Notion-Version": "2022-06-28",
                "Content-Type": "application/json",
            }
            cursor: Optional[str] = None
            async with httpx.AsyncClient(timeout=60.0) as client:
                while True:
                    body: Dict[str, Any] = {"page_size": 100}
                    if cursor:
                        body["start_cursor"] = cursor
                    r = await client.post(
                        f"https://api.notion.com/v1/databases/{TABLE_OF_CONTENT_DS}/query",
                        headers=headers,
                        json=body,
                    )
                    if r.status_code != 200:
                        logger.warning("ToC: REST query failed: %s %s", r.status_code, r.text[:200])
                        break
                    data = r.json()
                    rows.extend(data.get("results", []))
                    cursor = data.get("next_cursor")
                    if not cursor:
                        break

            if rows:
                logger.info("ToC: queried %d rows via REST API", len(rows))
        except Exception as e:
            logger.warning("ToC: REST query failed: %s", e)

    if not rows:
        logger.warning("ToC: no rows returned from database")
        return []

    # ── Fetch all rows ───────────────────────────────────────────────────
    import asyncio
    sem = asyncio.Semaphore(5)  # Limit concurrent fetches

    async def _fetch_with_limit(row: Dict) -> Optional[Dict]:
        async with sem:
            try:
                return await fetch_toc_row(row, google_creds=google_creds, use_cache=use_cache)
            except Exception as e:
                logger.debug("ToC: fetch failed for row: %s", e)
                return None

    results = await asyncio.gather(*[_fetch_with_limit(r) for r in rows])
    docs = [r for r in results if r is not None]

    # Stats
    by_type = {}
    for d in docs:
        ct = d.get("content_type", "unknown")
        by_type[ct] = by_type.get(ct, 0) + 1
    main_count = sum(1 for d in docs if d.get("priority") == "high")

    logger.info(
        "ToC: fetched %d/%d documents (%d main-source) — %s",
        len(docs), len(rows), main_count,
        ", ".join(f"{k}: {v}" for k, v in by_type.items()),
    )
    return docs
